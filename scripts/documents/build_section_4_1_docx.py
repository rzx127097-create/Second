"""Build the standalone Word deliverable for thesis section 4.1.

Equations are inserted as editable Office Math (OMML) by converting LaTeX
through Pandoc, then placed in an invisible two-column equation-number table.
"""

from __future__ import annotations

import copy
import re
import shutil
import subprocess
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Cm, Mm, Pt, RGBColor
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[2]
SOURCE_MD = ROOT / "docs" / "thesis" / "section-4.1.md"
FIG_DIR = ROOT / "artifacts" / "figures" / "chapter4"
OUT_REPO = ROOT / "artifacts" / "documents"
OUT_USER = Path(r"C:\Users\RZX\Desktop\论文\小论文\第二个问题\第二问")
OUTPUT_NAME = "4.1问题描述与空地协同保障机制设计.docx"

BODY_FONT_CN = "宋体"
BODY_FONT_EN = "Times New Roman"
HEADING_FONT_CN = "黑体"
MATH_FONT = "Cambria Math"


def set_east_asia(rpr, font: str) -> None:
    rfonts = rpr.rFonts
    rfonts.set(qn("w:eastAsia"), font)


def set_run_font(run, *, cn=BODY_FONT_CN, en=BODY_FONT_EN, size=Pt(12), bold=False) -> None:
    run.font.name = en
    run.font.size = size
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)
    set_east_asia(run._element.get_or_add_rPr(), cn)


def set_cell_margins(cell, top=45, start=45, bottom=45, end=45) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders_none(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            borders.append(node)
        node.set(qn("w:val"), "nil")


def set_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    flag = OxmlElement("w:tblHeader")
    flag.set(qn("w:val"), "true")
    tr_pr.append(flag)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.5)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = BODY_FONT_EN
    normal.font.size = Pt(12)
    set_east_asia(normal.element.get_or_add_rPr(), BODY_FONT_CN)
    pf = normal.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.first_line_indent = Cm(0.85)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.widow_control = True

    heading_specs = {
        "Heading 1": (16, 14, 8),
        "Heading 2": (14, 12, 6),
        "Heading 3": (12, 10, 4),
    }
    for name, (size, before, after) in heading_specs.items():
        style = styles[name]
        style.font.name = BODY_FONT_EN
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        set_east_asia(style.element.get_or_add_rPr(), HEADING_FONT_CN)
        fmt = style.paragraph_format
        fmt.alignment = WD_ALIGN_PARAGRAPH.LEFT
        fmt.first_line_indent = Cm(0)
        fmt.space_before = Pt(before)
        fmt.space_after = Pt(after)
        fmt.line_spacing = 1.0
        fmt.keep_with_next = True
        fmt.keep_together = True

    # Footer page field, centered and intentionally simple.
    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    footer_p._p.append(fld)

    settings = doc.settings._element
    compat = settings.find(qn("w:compat"))
    if compat is None:
        compat = OxmlElement("w:compat")
        settings.append(compat)
    mode = OxmlElement("w:compatSetting")
    mode.set(qn("w:name"), "compatibilityMode")
    mode.set(qn("w:uri"), "http://schemas.microsoft.com/office/word")
    mode.set(qn("w:val"), "15")
    compat.append(mode)


def add_text_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(0.85)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run)


def add_markdown_paragraph(doc: Document, markdown: str) -> None:
    """Convert one prose paragraph so every inline formula remains editable."""
    with tempfile.TemporaryDirectory(prefix="section41_para_") as tmp:
        tmp_dir = Path(tmp)
        md_path = tmp_dir / "paragraph.md"
        docx_path = tmp_dir / "paragraph.docx"
        md_path.write_text(markdown + "\n", encoding="utf-8")
        subprocess.run(
            ["pandoc", "--from", "markdown+tex_math_dollars", str(md_path), "-o", str(docx_path)],
            check=True,
            capture_output=True,
        )
        source = Document(docx_path)
        if not source.paragraphs:
            raise RuntimeError(f"Pandoc did not produce a paragraph for: {markdown[:80]}")
        element = copy.deepcopy(source.paragraphs[0]._p)

    body = doc._body._element
    sect_pr = body.find(qn("w:sectPr"))
    if sect_pr is None:
        body.append(element)
    else:
        body.insert(body.index(sect_pr), element)
    paragraph = doc.paragraphs[-1]
    paragraph.style = doc.styles["Normal"]
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Cm(0.85)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    for run in paragraph.runs:
        set_run_font(run)
    set_math_fonts(element)


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    size = {1: Pt(16), 2: Pt(14), 3: Pt(12)}[level]
    set_run_font(run, cn=HEADING_FONT_CN, size=size, bold=True)


def latex_to_omml(latex: str):
    with tempfile.TemporaryDirectory(prefix="section41_eq_") as tmp:
        tmp_dir = Path(tmp)
        md = tmp_dir / "eq.md"
        docx = tmp_dir / "eq.docx"
        md.write_text(f"$$\n{latex}\n$$\n", encoding="utf-8")
        subprocess.run(["pandoc", str(md), "-o", str(docx)], check=True, capture_output=True)
        eq_doc = Document(docx)
        nodes = eq_doc._element.xpath("//m:oMathPara|//m:oMath")
        if not nodes:
            raise RuntimeError(f"Pandoc did not produce OMML for: {latex}")
        node = copy.deepcopy(nodes[0])
        # Equation alignment is handled by the containing table cell.
        if node.tag == qn("m:oMathPara"):
            for pr in node.findall(qn("m:oMathParaPr")):
                node.remove(pr)
        return node


def set_math_fonts(omml) -> None:
    for rpr in omml.iter(qn("m:rPr")):
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.insert(0, rfonts)
        for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
            rfonts.set(qn(f"w:{attr}"), MATH_FONT)


def add_equation(doc: Document, latex: str, number: str) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders_none(table)
    table.columns[0].width = Cm(13.1)
    table.columns[1].width = Cm(1.8)
    table.rows[0].cells[0].width = Cm(13.1)
    table.rows[0].cells[1].width = Cm(1.8)
    for cell in table.rows[0].cells:
        set_cell_margins(cell, top=35, start=20, bottom=35, end=20)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    left = table.cell(0, 0).paragraphs[0]
    left.alignment = WD_ALIGN_PARAGRAPH.CENTER
    left.paragraph_format.first_line_indent = Cm(0)
    left.paragraph_format.space_before = Pt(2)
    left.paragraph_format.space_after = Pt(2)
    omml = latex_to_omml(latex)
    set_math_fonts(omml)
    left._p.append(omml)

    right = table.cell(0, 1).paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right.paragraph_format.first_line_indent = Cm(0)
    right.paragraph_format.space_before = Pt(2)
    right.paragraph_format.space_after = Pt(2)
    run = right.add_run(f"（{number}）")
    set_run_font(run, size=Pt(12))


def add_figure(doc: Document, image_path: Path, caption: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(str(image_path), width=Cm(15.5))

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.first_line_indent = Cm(0)
    cap.paragraph_format.line_spacing = 1.0
    cap.paragraph_format.space_before = Pt(0)
    cap.paragraph_format.space_after = Pt(6)
    cap.paragraph_format.keep_with_next = False
    run = cap.add_run(caption)
    set_run_font(run, size=Pt(10.5))


CONTENT = [
    ("h1", "4.1 问题描述与空地协同保障机制设计"),
    ("p", "在动态虫害演化与多无人机协同施药模型的基础上，本章进一步考虑机载药液有限时的持续作业问题。与将无人机药液视为无限资源的规划情形不同，有限载荷使单架无人机难以在整个决策周期内连续施药。当剩余药液不足时，无人机不仅需要调整虫害区域的访问顺序，还必须确定何时提出补给需求、前往何处与地面保障单元接驳，以及如何在等待与绕行代价可控的条件下恢复作业。与此同时，移动药液补给车受道路拓扑、行驶速度、剩余库存和服务能力约束，无法直接到达任意农田栅格。由此，原有的空中协同路径规划转化为空中施药任务与地面移动保障相互耦合的联合决策问题。"),
    ("p", "本节从作业实体、补给需求、服务过程和耦合关系四个层面明确研究边界。需要说明的是，本节仅建立模型规格与后续验证口径，不预设移动保障或 SR-MAPPO 的性能结论；无人机载药容量、车辆库存、行驶速度、补给速率、服务时间和接驳半径等工程参数将在实验设置中依据设备资料、田间研究或专家确认范围完成标定并冻结。"),
    ("h2", "4.1.1 作业场景与系统构成"),
    ("p", "研究区域表示为二维农田区域 Ω⊂R²，其中 R 表示实数集合；区域按空间分辨率离散为 H×W 个栅格。各栅格包含虫害密度、作物或地表属性以及药效状态等信息；虫害在生长、扩散、风场迁移和药剂作用下随时间演化。为兼顾生态过程的数值计算精度与智能体决策频率，本文区分生态积分步长 δt_eco 和物理决策步长 Δt_dec。一个决策步内可执行 K 次生态状态更新，即 Δt_dec=Kδt_eco，其中 K 为正整数。后续涉及飞行、车辆行驶、等待和服务的时间量均以物理时间计量，避免将仿真迭代次数直接等同于实际作业时间。"),
    ("p", "设无人机集合为 U={1,2,…,N_u}。无人机在农田栅格内执行移动、施药、等待和接驳等行为，并依据局部虫害状态、自身位置、剩余药液及可获得的协同信息进行决策。对无人机 i，以 x^u_i,t 表示时刻 t 的位置，以 q_i,t 和 q_i^max 分别表示其剩余药液量和载药容量，单位均为 L。药液只在有效施药或补给事件中发生变化，飞行能耗不在本章作为补给对象，因此不能将药液补给机制解释为电池补能机制。"),
    ("p", "地面保障系统由移动药液补给车和道路网络构成。道路网络抽象为带权图 G^r=(V^r,E^r)，其中 V^r 为道路节点集合，E^r 为可通行路段集合，边权用于表示道路距离或行驶时间。车辆集合记为 V={1,2,…,N_v}，车辆 v 在时刻 t 的道路位置和剩余库存分别记为 x^g_v,t 与 Q_v,t。本文主实验采用单辆补给车，以突出移动保障位置对协同施药的影响；集合化符号仅用于保持模型的可扩展性。车辆不能脱离道路图自由穿越农田，其策略层负责选择待服务请求或候选接驳点，确定性最短路执行器负责将该高层决策转换为道路上的可行路径。"),
    ("p", "候选接驳点集合记为 P_t。每个接驳点均与可达道路节点关联，并满足无人机能够进入服务半径的空间条件。接驳点既不是任意农田位置，也不是仅用于绘图的道路投影点；其有效性需要同时通过道路连通性、车辆可达性、无人机可达性和服务几何约束检验。图4-1给出了空地协同施药系统的实体构成与主要交互关系。其中，实线表示无人机或车辆的物理运动与药液转移，虚线表示补给请求、预计到达时间和库存等信息交互。道路数据仅用于构造具有实际拓扑特征的仿真输入，不作为真实部署验证的替代。"),
    ("fig", "fig4-1_air_ground_system.png", "图4-1  空地协同施药系统构成与物理交互"),
    ("h2", "4.1.2 有限机载药液下的动态补给需求"),
    ("p", "有限载药条件下，补给需求不应仅由固定药量阈值触发。原因在于，相同剩余药量对不同作业状态具有不同含义：处于高虫害热点并持续施药的无人机，其药液消耗速度较快；处于转场或低负荷状态的无人机，即使剩余药量相同，也可能具有更长的可持续作业时间。此外，补给车与无人机到候选接驳点的距离、道路绕行、已有服务队列和补给耗时都会影响请求的合理提前量。因此，本文以剩余药液可支持的作业时间作为需求评估基础。"),
    ("p", "设平均药液消耗率为根据近期施药记录和待执行任务估计的量，单位为 L/s；ε 为避免分母为零的正数。无人机 i 在时刻 t 的剩余药液可支持作业时间定义为"),
    ("eq", r"\hat T^{\mathrm{remain}}_{i,t}=\frac{q_{i,t}}{\max\!\left(\bar\rho^{\mathrm{spr}}_{i,t},\varepsilon\right)}", "4-1"),
    ("p", "式中，T 的估计值表示剩余药液可支持的作业时间，单位为 s；q_i,t 为无人机剩余药液量，平均药液消耗率的单位为 L/s。当无人机没有待执行施药任务或估计消耗率为零时，不依据式（4-1）单独生成紧急请求，而是保留其资源状态并在后续决策步重新评估。"),
    ("p", "设预计接驳时间为无人机从当前状态到可获得补给所需的时间，预计服务时间为完成该次服务所需的时间，ΔT_safe 为用于吸收虫害变化、道路绕行和到达时间估计误差的安全余量，单位均为 s。当满足"),
    ("eq", r"\hat T^{\mathrm{remain}}_{i,t}\leq\hat T^{\mathrm{rv}}_{i,t}+T^{\mathrm{svc}}_{i,t}+\Delta T^{\mathrm{safe}}", "4-2"),
    ("p", "时，系统允许无人机 i 生成补给请求。该条件的实质是比较“当前资源还能维持多久”与“获得下一次可用服务需要多久”，因而能够随无人机任务负荷、车机相对位置和服务拥堵程度动态变化。安全余量的取值将在敏感性实验中单独考察，避免通过过大的提前量人为降低等待风险。"),
    ("p", "每个请求至少记录请求编号、无人机编号、生成时刻、当前位置、剩余药量、目标补给量、紧迫度和可行候选接驳点。请求生成后进入开放状态，尚未被车辆接受时允许继续更新预计到达时间与紧迫度；一旦被某车辆预约，则建立请求—车辆—接驳点的唯一关联，以避免同一需求被重复服务。若无人机在等待过程中因任务变化、候选点失效或回合结束而取消请求，状态机必须记录取消原因，而不能将空请求误判为服务完成。"),
    ("p", "动态需求机制同时用于区分两类资源瓶颈。若系统总药液量不足，即使车辆能够即时到达，仍会出现大量未满足需求；若总药液基本充足，但库存长期停留在远离需求的位置，则主要表现为接驳距离、等待时间和药液失能时间增加。后续实验将通过无限药液、有限药液无补给、固定保障、即时补给诊断上界和道路约束移动保障等模式区分“总量不足”与“位置—时间错配”，从而检验移动补给机制是否真正被激活。"),
    ("h2", "4.1.3 车机接驳与补给服务机制"),
    ("p", "补给请求形成后，车辆策略从开放请求及其候选接驳点中选择服务目标。该决策属于任务层调度，而非直接输出车辆在栅格上的逐步移动方向。接驳点确定后，车辆沿道路图 G^r 的最短可行路径行驶；无人机则在保证施药与资源安全的前提下前往同一接驳区域。该分层设计使强化学习负责具有协同性和长期影响的目标选择，同时由确定性路径执行器保证车辆运动始终满足道路拓扑约束。"),
    ("p", "对请求无人机 i、车辆 v 和候选接驳点 p，设 d_u(x^u_i,t,p) 为无人机到接驳点的可行飞行距离，d_G(x^g_v,t,p) 为道路图上的最短路径距离，单位均为 m；v^u_i 和 v^g_v 分别为无人机与车辆的作业速度，单位为 m/s。忽略尚未发生的随机扰动时，二者预计到达时间及联合到达时间为"),
    ("eq", r"\hat T^u_{i,p,t}=\frac{d_u(\boldsymbol x^u_{i,t},p)}{v_i^u},\quad\hat T^g_{v,p,t}=\frac{d_{\mathcal G}(\boldsymbol x^g_{v,t},p)}{v_v^g},\quad\hat T^{\mathrm{arr}}_{i,v,p,t}=\max\!\left\{\hat T^u_{i,p,t},\hat T^g_{v,p,t}\right\}", "4-3"),
    ("p", "式（4-3）表明，服务开始时间由较晚到达的一方决定。设车辆处理已预约请求所需的预计排队时间为 T_queue，则无人机等待时间和车辆等待时间分别可表示为"),
    ("eq", r"\begin{aligned}\hat T^{\mathrm{wait},u}_{i,v,p,t}&=\max\!\left(0,\hat T^g_{v,p,t}-\hat T^u_{i,p,t}\right)+T^{\mathrm{queue}}_{v,t},\\[2pt]\hat T^{\mathrm{wait},g}_{i,v,p,t}&=\max\!\left(0,\hat T^u_{i,p,t}-\hat T^g_{v,p,t}\right).\end{aligned}", "4-4"),
    ("p", "其中，两类等待时间的单位均为 s。式（4-4）用于描述车机到达不同步造成的时间损失；实际请求等待时间还应从请求生成时刻持续累计至服务开始时刻，并在事件日志中保留请求排队、在途和接驳等待的分项记录。"),
    ("p", "只有当车辆到达对应道路节点、无人机进入接驳半径、请求归属一致、车辆库存大于零且双方均处于可服务状态时，服务才能开始。开始服务后，请求进入显式的服务锁定状态。在锁定期间，无人机不得执行移动或施药动作，车辆不得切换服务目标；可行动作集合及其掩码必须与该状态同步更新，并在策略采样和训练更新阶段保持一致。这样可避免环境在策略输出之后静默覆盖动作，进而造成行为策略与训练策略不一致。"),
    ("p", "设 q^svc_max 为单次服务允许转移的最大药液量，单位为 L。车辆 v 向无人机 i 的实际转移量定义为"),
    ("eq", r"\Delta q_{i,v,t}=\min\!\left\{q_i^{\max}-q_{i,t},\;Q_{v,t},\;q^{\mathrm{svc}}_{\max}\right\}", "4-5"),
    ("p", "式（4-5）同时限制无人机可用容量、车辆剩余库存和单次服务能力。当车辆库存不足以满足目标补给量时，系统执行部分补给并记录未满足量，而不将其误记为完整补给。若补给速率为 ρ_svc（L/s）、服务准备时间为 T_setup（s），则对应服务持续时间为 T_setup+Δq_i,v,t/ρ_svc。服务完成后释放双方锁定，无人机恢复施药决策，车辆重新进入可调度状态。车辆库存耗尽本身不触发回合终止；其后车辆不再具有有效补给动作，但无人机仍可利用剩余药液继续作业。"),
    ("p", "图4-2将上述过程组织为“时间预测—请求开放—预约与路由—联合到达—服务锁定—药液转移—恢复施药”的离散事件流程。开放、预约、在途、等待、服务、完成、取消和失败均应具有明确的进入条件、状态更新和事件记录字段。该流程既作为强化学习环境的状态转移依据，也作为动态需求调度—滚动 A* 等传统保障方法的共同服务接口，从而保证不同车辆策略在相同物理约束下比较。"),
    ("fig", "fig4-2_service_process.png", "图4-2  动态补给请求与车机接驳服务流程"),
    ("h2", "4.1.4 空地协同的时间、空间与资源关系"),
    ("p", "空地协同施药首先表现为空间耦合。无人机可以在农田区域内直接接近虫害热点，而车辆只能沿道路网络移动，两类主体的可达空间并不一致。候选接驳点需要在道路可达位置与无人机可接受绕行范围之间取得平衡：接驳点过于靠近车辆可能增加无人机脱离施药区域的距离，过于靠近虫害热点则可能导致车辆沿路网大幅绕行甚至不可达。因此，接驳决策的有效距离不能仅以欧氏距离衡量，还需同时考虑无人机飞行距离、车辆路网距离以及道路与农田之间的服务几何关系。"),
    ("p", "其次，补给过程具有显著的时间耦合。无人机过早到达会增加等待和施药中断，车辆过早到达则会占用保障资源并延迟后续请求；任何一方过晚到达都可能使无人机在完成当前任务前耗尽可用药液。补给请求的提前量、车机预计到达时间、已有服务队列和服务持续时间共同决定了接驳是否及时。因而，车辆调度不能只按照请求生成顺序执行，无人机也不能仅按最短飞行距离选择接驳点，而应结合请求紧迫度与联合到达关系进行协调。"),
    ("p", "再次，药液转移构成资源耦合。设 s_i,t 为决策步 t 内无人机 i 的实际施药量，单位为 L，则服务事件前后的资源更新满足"),
    ("eq", r"q_{i,t+1}=q_{i,t}-s_{i,t}+\sum_{v\in\mathcal V}\Delta q_{i,v,t},\qquad Q_{v,t+1}=Q_{v,t}-\sum_{i\in\mathcal U}\Delta q_{i,v,t}", "4-6"),
    ("p", "在忽略泄漏并允许数值舍入误差的条件下，系统总药液应满足"),
    ("eq", r"\sum_{i\in\mathcal U}q_{i,t}+\sum_{v\in\mathcal V}Q_{v,t}+\sum_{\tau=0}^{t-1}\sum_{i\in\mathcal U}s_{i,\tau}=\sum_{i\in\mathcal U}q_{i,0}+\sum_{v\in\mathcal V}Q_{v,0}", "4-7"),
    ("p", "式（4-6）和式（4-7）分别给出单步库存更新与全局药液守恒关系。任何补给量都必须在无人机端增加并在车辆端等量扣减；部分补给、库存耗尽和同时请求等情形均应通过同一守恒审计。该约束能够识别重复入账、负库存和环境状态更新遗漏，是后续算法训练前必须通过的确定性验证条件。"),
    ("p", "上述三类关系进一步产生任务层面的耦合：车辆服务某一请求，会改变其他无人机可获得服务的时间；无人机选择接驳点，会同时改变其施药覆盖、药液消耗和车辆后续位置；车辆完成一次补给后所在道路节点，又会影响下一请求的服务代价。因此，移动保障的潜在作用并非简单增加系统药液总量，而是调整药液资源在时间和空间上的分布。后续机制分析将依次考察移动保障是否缩短接驳距离、减少等待与药液失能时间、增加有效施药时间，并据此检验治理效果是否改善。"),
    ("h2", "4.1.5 研究目标与关键问题"),
    ("p", "针对上述空地异构协同过程，本文的研究目标是在虫害动态演化、无人机机载药液有限、车辆库存有限和道路可达性约束下，联合优化多无人机施药行为与移动补给车服务决策。设 η_red 为回合结束时的虫害消减率，I(η_red≥0.85) 为治理达标指示量，T_wait、T_off、D^u_rv 和 D^g 分别表示累计补给等待时间、药液失能时间、无人机接驳绕行距离和车辆道路行驶距离，则有限时域内的高层目标可概括为"),
    ("eq", r"\max_{\Pi}\;\mathbb E_{\Pi}\!\left[w_1\eta_{\mathrm{red}}+w_2\mathbb I\!\left(\eta_{\mathrm{red}}\geq0.85\right)-w_3T^{\mathrm{wait}}-w_4T^{\mathrm{off}}-w_5D^u_{\mathrm{rv}}-w_6D^g\right],\quad w_k\geq0", "4-8"),
    ("p", "其中，Π 表示空地异构主体的联合策略，w_k 为各目标项权重。式（4-8）用于说明治理收益、施药连续性和保障代价之间的总体关系，具体状态、动作、团队奖励和 SR-MAPPO 优化形式将在后续章节中给出。权重不在本节预先赋值，其设置需结合量纲处理、消融和敏感性分析进行确定。"),
    ("p", "围绕该目标，需要解决三个关键问题。第一，如何在无人机负责移动与施药、车辆负责需求服务选择且二者动作语义不同的条件下，建立共享治理目标下的空地异构联合决策机制。第二，如何将剩余作业时间、请求紧迫度、车机预计到达时间、道路可达性和车辆库存统一映射为可执行的请求—接驳点匹配决策。第三，如何通过显式服务状态、动作掩码、部分补给和药液守恒约束维持施药—接驳—补给—复作过程的一致性，并识别移动保障对治理结果的实际作用路径。"),
    ("p", "为此，后续研究将依次建立道路与资源约束下的空地异构决策模型，设计动态补给需求评估和候选接驳点生成机制，并在 SR-MAPPO 框架内实现角色独立策略与集中价值评估。实验部分则通过资源约束激活诊断、同资源条件下的固定与移动保障比较、强传统规划基线、规模适应性、参数敏感性和机制指标分析，检验所提出方法在不同场景中的适用范围。"),
]


def populate_from_markdown(doc: Document) -> None:
    lines = SOURCE_MD.read_text(encoding="utf-8").splitlines()
    i = 0
    skip_caption = False
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue
        if line.startswith("# "):
            add_heading(doc, line[2:].strip(), 1)
            i += 1
            continue
        if line.startswith("## "):
            add_heading(doc, line[3:].strip(), 2)
            i += 1
            continue
        if line.startswith("!["):
            match = re.match(r"!\[(.+?)\]\((.+?)\)", line)
            if not match:
                raise RuntimeError(f"Invalid figure markup: {line}")
            caption, rel_path = match.groups()
            add_figure(doc, FIG_DIR / Path(rel_path).name, caption)
            skip_caption = True
            i += 1
            continue
        if skip_caption and line.startswith("**图") and line.endswith("**"):
            skip_caption = False
            i += 1
            continue
        skip_caption = False
        if line == "$$":
            equation_lines = []
            i += 1
            while i < len(lines) and lines[i].strip() != "$$":
                equation_lines.append(lines[i].strip())
                i += 1
            if i >= len(lines):
                raise RuntimeError("Unclosed display equation in section-4.1.md")
            latex = " ".join(equation_lines)
            number_match = re.search(r"\\tag\{([^}]+)\}", latex)
            if not number_match:
                raise RuntimeError(f"Equation has no tag: {latex}")
            number = number_match.group(1)
            latex = re.sub(r"\s*\\tag\{[^}]+\}\s*", "", latex).strip()
            latex = re.sub(r"\s*\.\s*$", "", latex)
            add_equation(doc, latex, number)
            i += 1
            continue
        add_markdown_paragraph(doc, line)
        i += 1


def build() -> tuple[Path, Path]:
    doc = Document()
    configure_document(doc)
    populate_from_markdown(doc)

    props = doc.core_properties
    props.title = "4.1 问题描述与空地协同保障机制设计"
    props.subject = "路网约束下多无人机与移动药液补给车空地协同施药"
    props.keywords = "SR-MAPPO; 空地异构协同; 移动药液补给; 道路约束"
    props.comments = "M1 模型规格稿；公式为可编辑 Office Math 对象。"

    OUT_REPO.mkdir(parents=True, exist_ok=True)
    OUT_USER.mkdir(parents=True, exist_ok=True)
    repo_path = OUT_REPO / OUTPUT_NAME
    user_path = OUT_USER / OUTPUT_NAME
    doc.save(repo_path)
    shutil.copy2(repo_path, user_path)
    return repo_path, user_path


if __name__ == "__main__":
    for path in build():
        print(path)
