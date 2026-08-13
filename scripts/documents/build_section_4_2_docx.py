"""Build the standalone Word deliverable for thesis Section 4.2.

The Markdown source is authoritative. Pandoc converts every inline and display
formula to editable Office Math (OMML); display equations are then placed in an
invisible two-column table so the expression is centred and its number is
right-aligned without drawing borders.
"""

from __future__ import annotations

import copy
import re
import shutil
import subprocess
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
SOURCE_MD = ROOT / "docs" / "thesis" / "section-4.2.md"
FIG_DIR = ROOT / "artifacts" / "figures" / "chapter4"
OUT_REPO = ROOT / "artifacts" / "documents"
OUT_USER = Path(r"C:\Users\RZX\Desktop\论文\小论文\第二个问题\第二问")
OUTPUT_NAME = "4.2空地异构协同施药决策模型.docx"

BODY_FONT_CN = "宋体"
BODY_FONT_EN = "Times New Roman"
HEADING_FONT_CN = "黑体"
MATH_FONT = "Cambria Math"


def set_east_asia(rpr, font: str) -> None:
    rfonts = rpr.rFonts
    rfonts.set(qn("w:eastAsia"), font)


def set_run_font(run, *, cn=BODY_FONT_CN, en=BODY_FONT_EN, size=Pt(12), bold=False) -> None:
    run.font.name = en
    run.font.size = size
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)
    set_east_asia(run._element.get_or_add_rPr(), cn)


def set_cell_margins(cell, top=45, start=45, bottom=45, end=45) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders_none(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            borders.append(node)
        node.set(qn("w:val"), "nil")


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.5)

    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT_EN
    normal.font.size = Pt(12)
    set_east_asia(normal.element.get_or_add_rPr(), BODY_FONT_CN)
    pf = normal.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.first_line_indent = Cm(0.85)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.widow_control = True

    heading_specs = {
        "Heading 1": (16, 14, 8),
        "Heading 2": (14, 12, 6),
        "Heading 3": (12, 10, 4),
    }
    for name, (size, before, after) in heading_specs.items():
        style = doc.styles[name]
        style.font.name = BODY_FONT_EN
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        set_east_asia(style.element.get_or_add_rPr(), HEADING_FONT_CN)
        fmt = style.paragraph_format
        fmt.alignment = WD_ALIGN_PARAGRAPH.LEFT
        fmt.first_line_indent = Cm(0)
        fmt.space_before = Pt(before)
        fmt.space_after = Pt(after)
        fmt.line_spacing = 1.0
        fmt.keep_with_next = True
        fmt.keep_together = True

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    footer._p.append(field)

    settings = doc.settings._element
    compat = settings.find(qn("w:compat"))
    if compat is None:
        compat = OxmlElement("w:compat")
        settings.append(compat)
    mode = OxmlElement("w:compatSetting")
    mode.set(qn("w:name"), "compatibilityMode")
    mode.set(qn("w:uri"), "http://schemas.microsoft.com/office/word")
    mode.set(qn("w:val"), "15")
    compat.append(mode)


def insert_body_element(doc: Document, element) -> None:
    body = doc._body._element
    sect_pr = body.find(qn("w:sectPr"))
    if sect_pr is None:
        body.append(element)
    else:
        body.insert(body.index(sect_pr), element)


def set_math_fonts(element) -> None:
    for rpr in element.iter(qn("m:rPr")):
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.insert(0, rfonts)
        for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
            rfonts.set(qn(f"w:{attr}"), MATH_FONT)


def pandoc_docx(markdown: str, prefix: str) -> Document:
    with tempfile.TemporaryDirectory(prefix=prefix) as tmp:
        tmp_dir = Path(tmp)
        md_path = tmp_dir / "source.md"
        docx_path = tmp_dir / "source.docx"
        md_path.write_text(markdown, encoding="utf-8")
        subprocess.run(
            [
                "pandoc",
                "--from",
                "markdown+tex_math_dollars",
                str(md_path),
                "-o",
                str(docx_path),
            ],
            check=True,
            capture_output=True,
        )
        return Document(docx_path)


def add_markdown_paragraph(doc: Document, markdown: str) -> None:
    source = pandoc_docx(markdown + "\n", "section42_para_")
    if not source.paragraphs:
        raise RuntimeError(f"Pandoc produced no paragraph for: {markdown[:80]}")
    element = copy.deepcopy(source.paragraphs[0]._p)
    set_math_fonts(element)
    insert_body_element(doc, element)
    paragraph = doc.paragraphs[-1]
    paragraph.style = doc.styles["Normal"]
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Cm(0.85)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    for run in paragraph.runs:
        set_run_font(run)


def add_heading(doc: Document, text: str, level: int) -> None:
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text)
    set_run_font(
        run,
        cn=HEADING_FONT_CN,
        size={1: Pt(16), 2: Pt(14), 3: Pt(12)}[level],
        bold=True,
    )


def latex_to_omml(latex: str):
    source = pandoc_docx(f"$$\n{latex}\n$$\n", "section42_eq_")
    nodes = source._element.xpath("//m:oMathPara|//m:oMath")
    if not nodes:
        raise RuntimeError(f"Pandoc produced no OMML for: {latex}")
    node = copy.deepcopy(nodes[0])
    if node.tag == qn("m:oMathPara"):
        for prop in node.findall(qn("m:oMathParaPr")):
            node.remove(prop)
    set_math_fonts(node)
    return node


def add_equation(doc: Document, latex: str, number: str) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders_none(table)
    widths = (Cm(13.1), Cm(1.8))
    for idx, width in enumerate(widths):
        table.columns[idx].width = width
        table.rows[0].cells[idx].width = width
    for cell in table.rows[0].cells:
        set_cell_margins(cell, top=35, start=20, bottom=35, end=20)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    formula_paragraph = table.cell(0, 0).paragraphs[0]
    formula_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    formula_paragraph.paragraph_format.first_line_indent = Cm(0)
    formula_paragraph.paragraph_format.space_before = Pt(2)
    formula_paragraph.paragraph_format.space_after = Pt(2)
    formula_paragraph._p.append(latex_to_omml(latex))

    number_paragraph = table.cell(0, 1).paragraphs[0]
    number_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    number_paragraph.paragraph_format.first_line_indent = Cm(0)
    number_paragraph.paragraph_format.space_before = Pt(2)
    number_paragraph.paragraph_format.space_after = Pt(2)
    set_run_font(number_paragraph.add_run(f"（{number}）"), size=Pt(12))


def add_figure(doc: Document, image_path: Path, caption: str) -> None:
    if not image_path.is_file():
        raise FileNotFoundError(image_path)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.keep_with_next = True
    paragraph.add_run().add_picture(str(image_path), width=Cm(14.5))

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.first_line_indent = Cm(0)
    cap.paragraph_format.line_spacing = 1.0
    cap.paragraph_format.space_before = Pt(0)
    cap.paragraph_format.space_after = Pt(6)
    set_run_font(cap.add_run(caption), size=Pt(10.5))


def populate_from_markdown(doc: Document) -> None:
    lines = SOURCE_MD.read_text(encoding="utf-8").splitlines()
    i = 0
    skip_caption = False
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue
        if line.startswith("# "):
            add_heading(doc, line[2:].strip(), 1)
            i += 1
            continue
        if line.startswith("## "):
            add_heading(doc, line[3:].strip(), 2)
            i += 1
            continue
        if line.startswith("!["):
            match = re.match(r"!\[(.+?)\]\((.+?)\)", line)
            if not match:
                raise RuntimeError(f"Invalid figure markup: {line}")
            caption, rel_path = match.groups()
            add_figure(doc, FIG_DIR / Path(rel_path).name, caption)
            skip_caption = True
            i += 1
            continue
        if skip_caption and line.startswith("**图") and line.endswith("**"):
            skip_caption = False
            i += 1
            continue
        skip_caption = False
        if line == "$$":
            equation_lines: list[str] = []
            i += 1
            while i < len(lines) and lines[i].strip() != "$$":
                equation_lines.append(lines[i].strip())
                i += 1
            if i >= len(lines):
                raise RuntimeError("Unclosed display equation in section-4.2.md")
            latex = " ".join(equation_lines)
            number_match = re.search(r"\\tag\{([^}]+)\}", latex)
            if not number_match:
                raise RuntimeError(f"Equation has no tag: {latex}")
            number = number_match.group(1)
            latex = re.sub(r"\s*\\tag\{[^}]+\}\s*", "", latex).strip()
            latex = re.sub(r"\s*[,.]\s*$", "", latex)
            add_equation(doc, latex, number)
            i += 1
            continue
        add_markdown_paragraph(doc, line)
        i += 1


def build() -> tuple[Path, Path]:
    doc = Document()
    configure_document(doc)
    populate_from_markdown(doc)

    props = doc.core_properties
    props.title = "4.2 空地异构协同施药决策模型"
    props.subject = "路网约束下多无人机与移动药液补给车的空地异构协同建模"
    props.keywords = "SR-MAPPO; 空地异构协同; 移动药液补给; Dec-POMDP; 动作掩码"
    props.comments = "M1模型规格正文；公式为可编辑Office Math对象。"

    OUT_REPO.mkdir(parents=True, exist_ok=True)
    OUT_USER.mkdir(parents=True, exist_ok=True)
    repo_path = OUT_REPO / OUTPUT_NAME
    user_path = OUT_USER / OUTPUT_NAME
    doc.save(repo_path)
    shutil.copy2(repo_path, user_path)
    return repo_path, user_path


if __name__ == "__main__":
    for output in build():
        print(output)
